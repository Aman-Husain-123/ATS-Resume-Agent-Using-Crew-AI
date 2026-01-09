"""
Utility Functions for Resume Processing
This module provides helper functions for converting and formatting resume text.
"""

from docx import Document
from io import BytesIO


def txt_to_docx_bytes(text: str) -> bytes:
    """
    Convert plain text to DOCX format in memory.
    
    This function creates a Word document from plain text by:
    - Splitting text into lines
    - Creating paragraphs for each line
    - Preserving empty lines for spacing
    - Returning document as bytes (no file I/O)
    
    Useful for:
    - Downloading optimized resumes as DOCX
    - Converting ATS-optimized text to formatted documents
    - Providing user-friendly output format
    
    Args:
        text: Plain text resume content (can include newlines)
        
    Returns:
        bytes: DOCX file content as bytes, ready for download or saving
        
    Example:
        >>> resume_text = "John Doe\\nSoftware Engineer\\n\\nExperience:\\n- Led team of 5"
        >>> docx_bytes = txt_to_docx_bytes(resume_text)
        >>> # Can now save or send docx_bytes as a file
    """
    # Create new Word document
    doc = Document()
    
    # Process each line of text
    for line in text.splitlines():
        if line.strip() == "":
            # Preserve empty lines for spacing
            doc.add_paragraph("")
        else:
            # Add line as paragraph
            doc.add_paragraph(line)
    
    # Save document to bytes buffer (in-memory)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
